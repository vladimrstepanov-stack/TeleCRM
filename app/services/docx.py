"""Формирование деловых DOCX-документов (редактируемых перед отправкой).

Важно: в документы НЕ попадают секретные данные — минимальная цена собственника
и зашифрованные персональные данные, а также внутренний жаргон бота.
"""

from datetime import datetime
from decimal import Decimal
from pathlib import Path

from docx import Document
from docx.shared import Pt

from app.database.models import Activity, AuditEvent
from app.database.repositories import Card

_ACTIVITY_TITLES = {
    "note": "Заметка",
    "call": "Звонок",
    "showing": "Показ",
    "negotiation": "Переговоры",
    "offer": "Оферта",
    "status_change": "Смена статуса",
    "price_change": "Изменение цены",
}
_STATUS_TITLES = {
    "active": "в продаже",
    "reserved": "бронь",
    "sold": "продан",
    "withdrawn": "снят с продажи",
    "new": "новая",
    "in_progress": "в работе",
    "offer": "оферта",
    "deposit": "задаток",
    "closed": "закрыта",
    "cancelled": "отменена",
}


def _money(value: Decimal | None) -> str:
    if value is None:
        return "не указана"
    return f"{value:,.0f} руб.".replace(",", " ")


def _status(value: str | None) -> str:
    return _STATUS_TITLES.get(value or "", value or "не указан")


def _property_title(card: Card) -> str:
    parts = [card.property.city, card.property.address]
    if card.property.house_number and card.property.house_number != "-":
        parts.append("д. " + card.property.house_number)
    if card.property.apartment_number:
        parts.append("кв. " + card.property.apartment_number)
    return ", ".join(p for p in parts if p and p != "не указан")


def _base_document(title: str) -> Document:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    heading = document.add_heading(title, level=0)
    heading.alignment = 1
    date_line = document.add_paragraph(datetime.now().strftime("Дата: %d.%m.%Y"))
    date_line.alignment = 1
    document.add_paragraph()
    return document


def _add_field(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(value)


def build_owner_report(
    path: Path, card: Card, activities: list[Activity], events: list[AuditEvent]
) -> Path:
    """Отчёт собственнику: актуальное сверху, хронология от новых к старым."""
    document = _base_document("Отчёт собственнику")

    document.add_heading("Актуальная информация", level=1)
    _add_field(document, "Клиент", card.client.name)
    _add_field(document, "Объект", _property_title(card))
    _add_field(document, "Текущая цена", _money(card.property.price))
    _add_field(document, "Статус объекта", _status(card.property.status))
    _add_field(document, "Статус работы", _status(card.deal.status))

    document.add_heading("Хронология работы", level=1)
    timeline = _build_timeline(activities, events)
    if not timeline:
        document.add_paragraph("Событий пока нет.")
    for moment, text in timeline:
        paragraph = document.add_paragraph(style="List Bullet")
        stamp = paragraph.add_run(moment.strftime("%d.%m.%Y %H:%M") + " — ")
        stamp.bold = True
        paragraph.add_run(text)

    document.save(path)
    return path


def _build_timeline(
    activities: list[Activity], events: list[AuditEvent]
) -> list[tuple[datetime, str]]:
    items: list[tuple[datetime, str]] = []

    for activity in activities:
        title = _ACTIVITY_TITLES.get(activity.activity_type, activity.activity_type)
        text = f"{title}. {activity.summarized_action}"
        if activity.buyer_feedback:
            text += f" Обратная связь покупателя: {activity.buyer_feedback}."
        if activity.seller_feedback:
            text += f" Обратная связь собственника: {activity.seller_feedback}."
        if activity.proposed_price is not None:
            text += f" Предложенная цена: {_money(activity.proposed_price)}."
        items.append((activity.created_at, text))

    for event in events:
        if event.event_type == "create":
            items.append((event.created_at, "Объект добавлен в работу."))
        elif event.event_type == "update" and event.changes_json:
            described = _describe_changes(event.changes_json)
            if described:
                items.append((event.created_at, described))
        elif event.event_type == "merge":
            items.append((event.created_at, "Карточки клиента объединены."))

    # Новые события сверху.
    items.sort(key=lambda pair: pair[0], reverse=True)
    return items


def _describe_changes(changes: dict) -> str:
    labels = {
        "price": "цена",
        "status": "статус объекта",
        "offer_price": "цена оферты",
    }
    parts = []
    for field, pair in changes.items():
        if not isinstance(pair, list) or len(pair) != 2:
            continue
        label = labels.get(field, field)
        old, new = pair
        if field in {"status"}:
            old, new = _status(old), _status(new)
        parts.append(f"{label}: {old} -> {new}")
    return "Изменения: " + "; ".join(parts) if parts else ""


def build_meeting_doc(path: Path, card: Card) -> Path:
    """Сводка «К встрече» для передачи клиенту/покупателю/продавцу."""
    document = _base_document("Сводка к встрече")

    document.add_heading("Клиент", level=1)
    _add_field(document, "Имя", card.client.name)

    document.add_heading("Объект", level=1)
    _add_field(document, "Адрес", _property_title(card))
    if card.property.property_type and card.property.property_type != "не указан":
        _add_field(document, "Тип", card.property.property_type)
    if card.property.rooms_count is not None:
        _add_field(document, "Комнат", str(card.property.rooms_count))
    if card.property.total_area is not None:
        _add_field(document, "Площадь", f"{card.property.total_area} кв. м")
    if card.property.floor is not None and card.property.total_floors is not None:
        _add_field(document, "Этаж", f"{card.property.floor} из {card.property.total_floors}")
    _add_field(document, "Цена", _money(card.property.price))

    document.add_heading("Статус", level=1)
    _add_field(document, "Объект", _status(card.property.status))
    _add_field(document, "Работа по сделке", _status(card.deal.status))
    if card.deal.notes:
        document.add_heading("Ключевые факты", level=1)
        document.add_paragraph(card.deal.notes)

    document.save(path)
    return path
